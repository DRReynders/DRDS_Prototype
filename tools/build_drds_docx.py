#!/usr/bin/env python3
"""
DRDS — Markdown -> DOCX export renderer.

PROMOTED TOOL. One renderer, one config per document type. Promoted from the
iSmile Dental CT rehearsal, where it built and validated two document shapes:
the Owner Report (baseline v4.6) and the Practitioner Brief (baseline v2.5).

    NOTE ON LANGUAGE: this is a promoted rehearsal tool, not a template engine.
    It renders a deliberately small Markdown subset that the DRDS report design
    system uses. Do not generalise it speculatively — extend it when a real
    document needs something, and test that document.

SOURCE-OF-TRUTH RULE
    The Markdown is the source of truth. This script reads Markdown and writes a
    DOCX to the configured output path. It must never read a DOCX, never write
    back to the Markdown, and never be used to "fix" content. Content problems
    are fixed in the Markdown and re-exported.

USAGE
    py -3 tools/build_drds_docx.py --config <path-to-json>

    Requires python-docx. PDF conversion is a SEPARATE step (LibreOffice
    headless) and is deliberately not part of this tool.

    A config is effectively required: the embedded DEFAULTS carry style and
    behaviour only, and hold NO paths. A build with no root/source_md/out_docx
    aborts rather than guessing. See load_config() for why.

HISTORY (as the iSmile rehearsal renderer, build_owner_report_docx.py)
    v4.0  first test export (US Letter)
    v4.1  callout panels, horizontal-rule tables, rule suppression,
          configurable source comments, inline-formatting bug fix
    v4.3  A4 page size
    v2.1  explicit-config abort, table-row split protection
    v2.2  callout split protection
    v4.4/v2.3  shared cover cosmetic policy, blank-header tables, post-build check
    v4.5/v2.4  selective page-break experiment
    v4.6/v2.5  selective page-break policy adopted — current export baselines
    ---   22 pt cleanup: dead spacing assignment and its false rationale removed.
          Output proved byte-identical to the v4.6/v2.5 baselines at the
          word/document.xml level.
    ---   PROMOTED to the Engineering Workspace as build_drds_docx.py. Renderer
          logic unchanged; client-specific path defaults removed (see DEFAULTS).
"""
import re, os, sys, json, shutil
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# DEFAULTS — style and behaviour only. NO PATHS.
#
# The rehearsal copy of this script shipped a real client's root, source and
# output path as defaults. That is safe in a one-client folder and actively
# dangerous in shared tooling: a config typo would silently build one client's
# document over another's approved export. It nearly did exactly that once, which
# is why load_config() aborts on an unreadable explicit --config.
#
# So the paths below are empty and every build must supply them. Everything else
# — palette, type scale, page geometry, guardrail switches — keeps the proven
# defaults, so a config only has to state what it actually differs on.
# ---------------------------------------------------------------------------
DEFAULTS = {
    # --- paths (no defaults — must come from the config) -------------------
    "root": "",          # absolute path the three paths below are relative to
    "source_md": "",     # Markdown source; the source of truth
    "out_docx": "",      # generated DOCX
    "assets_dir": "",    # folder holding the logo
    "logo_file": "DRDS_Logo_Mark_OnWhite.png",
    "logo_width_in": 1.45,

    # --- page geometry -----------------------------------------------------
    # A4 is the DRDS default. Letter is US-only and should be a deliberate
    # override, never an accident. python-docx defaults to Letter if unset,
    # which is how v4.0-v4.2 shipped at the wrong size.
    "page_size": "A4",            # A4 | Letter
    "margin_top_in": 1.0, "margin_bottom_in": 1.0,
    "margin_left_in": 1.1, "margin_right_in": 1.1,

    # --- palette -----------------------------------------------------------
    "navy": "1B2A4A", "charcoal": "333333", "gold": "B08D57", "grey": "6B6B6B",
    "callout_fill": "F6F4F0", "callout_bar": "B08D57",
    "internal_fill": "FBF1F1", "internal_bar": "9B2C2C", "internal_text": "9B2C2C",
    "rule_colour": "D8D2C8", "table_rule": "C9C2B6", "table_head_fill": "EFECE6",

    # --- typography --------------------------------------------------------
    "head_font": "Cambria", "body_font": "Calibri", "mono_font": "Consolas",
    "body_pt": 10.5, "title_pt": 30, "h1_pt": 19, "h2_pt": 15, "h3_pt": 11.5,

    # --- behaviour ---------------------------------------------------------
    "page_break_on_h1": True,
    "suppress_rules_near_headings": True,
    "callout_as_panel": True,
    "table_style": "horizontal_only",
    "source_comments": "compact",        # full | compact | off
    "render_internal_markers": True,     # MUST stay true until client-ready

    # --- table pagination (added v2.1) ------------------------------------
    # A short table splitting across a page break orphans its last row, which
    # gives that row unearned prominence. Content tables at or under the row
    # threshold are kept whole; longer ones are allowed to split, because
    # forcing a long table onto one page creates a worse gap than it solves.
    # Callout panels are never affected — see keep_table_together().
    "table_rows_cant_split": True,
    "table_keep_together_max_rows": 8,

    # --- callout pagination (added v2.2) ----------------------------------
    # A callout is a single-row table. Without cantSplit it breaks mid-sentence
    # across a page boundary, which destroys the "one continuous panel" property
    # the v4.1 refinement exists to provide. cantSplit degrades gracefully: Word
    # and LibreOffice ignore it when a row genuinely cannot fit on one page, so
    # an over-tall callout still splits rather than vanishing.
    "callout_cant_split": True,

    # --- shared cosmetic policy (added v4.4 / v2.3) ------------------------
    # The "cover" is everything before the first page-breaking section heading.
    # Source comments are build metadata, NOT internal safety markers: they are
    # suppressed on the cover so the page opens on the logo, while the rehearsal
    # boundary, placeholders and brackets always render.
    # Selective page breaks. page_break_on_h1 sets the default for every section
    # heading; any heading whose text starts with one of these prefixes is
    # exempted and allowed to flow. Kept as data, not logic, so no document's
    # structure is hard-coded into the renderer. An exempted heading is treated
    # identically to any other section heading apart from the missing break — it
    # gains no extra spacing above.
    "no_page_break_h1_prefixes": [],
    "suppress_cover_source_comments": True,
    # Markdown --- rules on the cover read as floating lines once the shaded
    # panels already separate blocks. Body and heading rules are unaffected.
    "suppress_cover_rules": True,
    # A table written with an empty header row is a label/value table, not a
    # header table. Style the first column, do not shade the first data row.
    "blank_header_as_label_value": True,
    # Post-build self-check. Verifies what can be verified without pagination.
    "post_build_check": True,
}

PAGE_SIZES_MM = {"A4": (210, 297), "Letter": (215.9, 279.4)}

def load_config():
    """Load configuration.

    SAFETY RULE (added v2.1, after a real incident):
    If a config was named EXPLICITLY with --config and cannot be read, the build
    ABORTS. It must never fall back to embedded defaults: shared tooling carries
    no path defaults, so an explicitly requested config has to load successfully
    rather than leaving the renderer to guess its build context. Falling back is
    only safe when no config was requested at all.

    The incident behind the rule: in the rehearsal renderer the embedded defaults
    DID hold a real client's output path, and a silent fallback overwrote an
    approved Owner Report export exactly that way. The promoted tool removes
    those path defaults as well — see DEFAULTS — but the abort stays, because
    guessing a build context is the failure, not merely guessing a bad one.

    'utf-8-sig' is used deliberately: a BOM is common in JSON written on Windows
    and is not a reason to fail.

    The implicit sibling lookup is retained from the rehearsal tool but now looks
    for the generically named drds_docx_config.json. The workspace ships only
    *.example.json files, so nothing is picked up implicitly here; a real build
    names its config explicitly and gets the abort behaviour above.
    """
    cfg = dict(DEFAULTS)
    path, explicit = None, False
    if "--config" in sys.argv:
        path = sys.argv[sys.argv.index("--config") + 1]
        explicit = True
    else:
        sibling = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                               "drds_docx_config.json")
        if os.path.exists(sibling):
            path = sibling
    if path:
        try:
            with open(path, encoding="utf-8-sig") as fh:
                cfg.update(json.load(fh))
            print(f"CONFIG: merged from {path}")
        except Exception as exc:
            if explicit:
                raise SystemExit(
                    f"ABORTED: --config {path} could not be read ({exc}).\n"
                    f"Refusing to fall back to embedded defaults. Shared tooling "
                    f"carries no path defaults, so an explicitly requested config "
                    f"must load successfully rather than allowing the renderer to "
                    f"guess its build context. Fix the config and re-run."
                )
            print(f"CONFIG WARNING: could not read {path} ({exc}); using defaults")
    else:
        print("CONFIG: embedded defaults")

    # Added at promotion. The embedded defaults deliberately hold no paths, so a
    # build that reaches here without them has no document to render and no place
    # to put it. Fail with the missing keys named rather than raising an opaque
    # path error twenty lines later.
    missing = [k for k in ("root", "source_md", "out_docx", "assets_dir")
               if not str(cfg.get(k, "")).strip()]
    if missing:
        raise SystemExit(
            f"ABORTED: no path configuration. Missing: {', '.join(missing)}.\n"
            f"This renderer carries style defaults only — never paths — so it "
            f"cannot guess which document to build or where to write it.\n"
            f"Run it as: py -3 build_drds_docx.py --config <your-config.json>\n"
            f"Start from drds_owner_report_docx_config.example.json or "
            f"drds_practitioner_brief_docx_config.example.json."
        )
    return cfg

CFG = load_config()

# Guard rail: turning internal markers off is a client-ready action and must be
# deliberate. Refuse silently proceeding if the source still declares itself a
# rehearsal draft.
def guard_internal_markers(md_text):
    if not CFG["render_internal_markers"] and "INTERNAL REHEARSAL DRAFT" in md_text:
        raise SystemExit(
            "REFUSED: render_internal_markers is false but the source Markdown "
            "still contains the rehearsal boundary. Turning markers off is a "
            "client-ready action and must be a recorded decision."
        )

R = lambda h: RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
NAVY, CHARCOAL, GOLD, GREY = (R(CFG["navy"]), R(CFG["charcoal"]),
                              R(CFG["gold"]), R(CFG["grey"]))
INTERNAL_TXT = R(CFG["internal_text"])

ROOT = CFG["root"]
SRC = os.path.join(ROOT, CFG["source_md"])
OUT = os.path.join(ROOT, CFG["out_docx"])
ASSETS = os.path.join(ROOT, CFG["assets_dir"])
LOGO = os.path.join(ASSETS, CFG["logo_file"])
os.makedirs(os.path.dirname(OUT), exist_ok=True)
os.makedirs(ASSETS, exist_ok=True)

# ---------------------------------------------------------------------------
# OOXML helpers
# ---------------------------------------------------------------------------
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn('w:' + k), v)
    return e

def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(_el('w:shd', val='clear', color='auto', fill=fill))

def cell_borders(cell, spec):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        s = spec.get(edge)
        b.append(_el(f'w:{edge}', val='nil') if not s else
                 _el(f'w:{edge}', val='single', sz=s[0], space='0', color=s[1]))
    tcPr.append(b)

def cell_margins(cell, top=90, left=140, bottom=90, right=140):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for edge, v in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        m.append(_el(f'w:{edge}', w=str(v), type='dxa'))
    tcPr.append(m)

def table_borders(table, spec):
    b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        s = spec.get(edge)
        b.append(_el(f'w:{edge}', val='nil') if not s else
                 _el(f'w:{edge}', val='single', sz=s[0], space='0', color=s[1]))
    table._tbl.tblPr.append(b)

def para_border(par, edge, color, sz):
    pPr = par._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr'); pPr.append(pbdr)
    pbdr.append(_el(f'w:{edge}', val='single', sz=sz, space='4', color=color))

def keep_with_next(par):
    par._p.get_or_add_pPr().append(_el('w:keepNext', val='true'))

def row_cant_split(row):
    """Stop a single row's content breaking across a page boundary."""
    row._tr.get_or_add_trPr().append(_el('w:cantSplit', val='true'))

def keep_table_together(tb):
    """Keep a whole table on one page.

    Word/LibreOffice have no 'keep table together' property. The reliable
    equivalent is keepNext on every paragraph of every row EXCEPT the last: each
    row is then tied to the row below it, so the table moves whole rather than
    splitting. The last row is left alone deliberately — keepNext there would
    tie the table to whatever paragraph follows and drag unrelated content with
    it.

    Applied to CONTENT tables only, never to callout panels: a callout is a
    single-row table and can be tall, so forcing it whole could push a large
    block to the next page and leave a bigger gap than it solves.
    """
    rows = list(tb.rows)
    for r in rows[:-1]:
        for c in r.cells:
            for p in c.paragraphs:
                keep_with_next(p)

# ---------------------------------------------------------------------------
INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')

def add_runs(par, text, size=None, color=CHARCOAL, font=None, bold_all=False):
    size = size or CFG["body_pt"]; font = font or CFG["body_font"]
    for part in INLINE.split(text):
        if not part:
            continue
        r = par.add_run()
        if part.startswith('**') and part.endswith('**'):
            r.text = part[2:-2]; r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r.text = part[1:-1]; r.font.name = CFG["mono_font"]
            r.font.size = Pt(size - 1.5); r.font.color.rgb = GREY
            continue
        elif part.startswith('*') and part.endswith('*'):
            r.text = part[1:-1]; r.italic = True
        else:
            r.text = part
        r.font.name = font; r.font.size = Pt(size); r.font.color.rgb = color
        if bold_all:
            r.bold = True

# ---------------------------------------------------------------------------
md_text = open(SRC, encoding='utf-8').read()
guard_internal_markers(md_text)
lines = md_text.split('\n')

doc = Document()
sec = doc.sections[0]
w_mm, h_mm = PAGE_SIZES_MM.get(CFG["page_size"], PAGE_SIZES_MM["A4"])
sec.page_width = Mm(w_mm)
sec.page_height = Mm(h_mm)
sec.top_margin = Inches(CFG["margin_top_in"]); sec.bottom_margin = Inches(CFG["margin_bottom_in"])
sec.left_margin = Inches(CFG["margin_left_in"]); sec.right_margin = Inches(CFG["margin_right_in"])

n = doc.styles['Normal']
n.font.name = CFG["body_font"]; n.font.size = Pt(CFG["body_pt"]); n.font.color.rgb = CHARCOAL
n.paragraph_format.space_after = Pt(8); n.paragraph_format.line_spacing = 1.15

ST = dict(h1=0, h2=0, h3=0, tables=0, callouts=0, internal_callouts=0, bullets=0,
          images=0, page_breaks=0, rules_kept=0, rules_suppressed=0,
          comments_full=0, comments_compact=0, comments_off=0,
          tables_kept_together=0, tables_allowed_to_split=0,
          cover_comments_suppressed=0, cover_rules_suppressed=0,
          label_value_tables=0, header_tables=0, page_breaks_exempted=0)

# The cover runs from the top of the file to the first page-breaking section
# heading. Cover-only cosmetic rules apply while this is True.
IN_COVER = True

def is_heading(idx):
    return 0 <= idx < len(lines) and lines[idx].strip().startswith('#')

def next_content_idx(idx):
    j = idx
    while j < len(lines) and not lines[j].strip():
        j += 1
    return j

def prev_content_idx(idx):
    j = idx
    while j >= 0 and not lines[j].strip():
        j -= 1
    return j

def add_callout(block, internal):
    """Single-cell borderless table = one continuous panel (no paragraph seams)."""
    fill = CFG["internal_fill"] if internal else CFG["callout_fill"]
    bar = CFG["internal_bar"] if internal else CFG["callout_bar"]
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(t, {})
    if CFG.get("callout_cant_split", True):
        row_cant_split(t.rows[0])
    cell = t.cell(0, 0)
    shade_cell(cell, fill); cell_borders(cell, {'left': ('20', bar)}); cell_margins(cell)

    # Join wrapped source lines into logical paragraphs BEFORE inline parsing,
    # otherwise **bold** / *italic* spanning two source lines never closes.
    paras, buf = [], []
    for b in block:
        if not b.strip():
            if buf:
                paras.append(' '.join(buf)); buf = []
        elif b.startswith('### '):
            if buf:
                paras.append(' '.join(buf)); buf = []
            paras.append(b)
        else:
            buf.append(b.strip())
    if buf:
        paras.append(' '.join(buf))

    first = True
    for b in paras:
        p = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        p.paragraph_format.space_after = Pt(5); p.paragraph_format.space_before = Pt(0)
        if b.startswith('### '):
            add_runs(p, b[4:], size=CFG["h3_pt"],
                     color=INTERNAL_TXT if internal else NAVY,
                     font=CFG["head_font"], bold_all=True)
        else:
            add_runs(p, b, color=INTERNAL_TXT if internal else CHARCOAL)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(0)
    ST["internal_callouts" if internal else "callouts"] += 1
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_comment(buf):
    mode = CFG["source_comments"]
    if IN_COVER and CFG.get("suppress_cover_source_comments", True):
        # Build metadata must not be the first thing on the cover. It stays in
        # the Markdown and the QA notes; it just does not open the document.
        ST["cover_comments_suppressed"] += 1
        return
    if mode == 'off':
        ST["comments_off"] += 1; return
    body = [b for b in buf if b.strip()]
    if not body:
        return
    if mode == 'compact':
        ST["comments_compact"] += 1
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"[INTERNAL SOURCE COMMENT — {len(body)} lines — {body[0][:110]}… "
                      f"Full text in the Markdown source; removed at client-ready.]")
        r.font.name = CFG["mono_font"]; r.font.size = Pt(7.5)
        r.font.color.rgb = GREY; r.italic = True
        return
    ST["comments_full"] += 1
    add_callout(body, internal=True)

# ---------------------------------------------------------------------------
i = 0; first_h1 = True; in_comment = False; cbuf = []

while i < len(lines):
    t = lines[i].strip()

    if t.startswith('<!--'):
        in_comment = True; cbuf = [t.replace('<!--', '').strip()]
        if t.endswith('-->'):
            in_comment = False; cbuf[0] = cbuf[0].replace('-->', '').strip()
            add_comment(cbuf); cbuf = []
        i += 1; continue
    if in_comment:
        if t.endswith('-->'):
            cbuf.append(t.replace('-->', '').strip()); in_comment = False
            add_comment(cbuf); cbuf = []
        else:
            cbuf.append(t)
        i += 1; continue

    if not t:
        i += 1; continue

    if re.match(r'^!\[(.*?)\]\((.*?)\)$', t):
        p = doc.add_paragraph()
        p.add_run().add_picture(LOGO, width=Inches(CFG["logo_width_in"]))
        p.paragraph_format.space_after = Pt(16)
        ST["images"] += 1; i += 1; continue

    if re.match(r'^-{3,}$', t):
        j = i
        while j < len(lines) and re.match(r'^-{3,}$', lines[j].strip()):
            j += 1
        near = is_heading(next_content_idx(j)) or is_heading(prev_content_idx(i - 1))
        if IN_COVER and CFG.get("suppress_cover_rules", True):
            ST["cover_rules_suppressed"] += 1
        elif CFG["suppress_rules_near_headings"] and near:
            ST["rules_suppressed"] += 1
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
            para_border(p, 'bottom', CFG["rule_colour"], '6')
            ST["rules_kept"] += 1
        i = j; continue

    if t.startswith('>'):
        block = []
        while i < len(lines) and lines[i].strip().startswith('>'):
            block.append(re.sub(r'^>\s?', '', lines[i].strip())); i += 1
        joined = "\n".join(block)
        internal = ('INTERNAL REHEARSAL DRAFT' in joined or 'FOUNDER' in joined
                    or 'PRODUCT COUNCIL' in joined or 'INTERNAL —' in joined)
        add_callout(block, internal); continue

    if t.startswith('|'):
        # Collect the raw block first, then split it on the alignment row.
        # A blank header (`| | |`) and the alignment row (`|---|---|`) both look
        # like "no content", so they are distinguished by the presence of a dash:
        # only the alignment row contains one. Without this, the blank header was
        # discarded as a separator and the FIRST DATA ROW inherited header
        # styling — the defect this fixes.
        raw_block = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            raw_block.append(lines[i].strip().strip('|')); i += 1
        sep_at = next((k for k, r in enumerate(raw_block)
                       if re.match(r'^[\s:\-|]+$', r) and '-' in r), None)
        if sep_at is None:
            head_raw, body_raw = [], raw_block
        else:
            head_raw, body_raw = raw_block[:sep_at], raw_block[sep_at + 1:]
        split_cells = lambda r: [c.strip() for c in r.split('|')]
        head_cells = [split_cells(r) for r in head_raw]
        blank_header = (not head_cells) or all(
            all(c == '' for c in hc) for hc in head_cells)
        label_value = blank_header and CFG.get("blank_header_as_label_value", True)
        rows = ([] if blank_header else head_cells) + [split_cells(r) for r in body_raw]
        ST["label_value_tables" if label_value else "header_tables"] += 1
        if rows:
            ncols = max(len(r) for r in rows)
            tb = doc.add_table(rows=0, cols=ncols); tb.alignment = WD_TABLE_ALIGNMENT.LEFT
            table_borders(tb, {'top': ('6', CFG["table_rule"]),
                               'bottom': ('6', CFG["table_rule"]),
                               'insideH': ('4', CFG["table_rule"])})
            for ridx, row in enumerate(rows):
                cells = tb.add_row().cells
                for cidx in range(ncols):
                    cell = cells[cidx]
                    cell_margins(cell, top=70, left=90, bottom=70, right=90)
                    cp = cell.paragraphs[0]
                    cp.paragraph_format.space_after = Pt(2)
                    cp.paragraph_format.space_before = Pt(2)
                    # A real header row only exists when the Markdown supplied
                    # header text. Label/value tables style column one instead,
                    # and never shade a data row.
                    head = (ridx == 0 and not label_value)
                    label = (cidx == 0 and label_value)
                    add_runs(cp, row[cidx] if cidx < len(row) else '', size=9.5,
                             color=NAVY if (head or label) else CHARCOAL,
                             bold_all=head or label)
                    if head:
                        shade_cell(cell, CFG["table_head_fill"])
            # --- keep-together handling (content tables only) ---------------
            if CFG.get("table_rows_cant_split", True):
                for r_ in tb.rows:
                    row_cant_split(r_)
            if len(rows) <= CFG.get("table_keep_together_max_rows", 8):
                keep_table_together(tb)
                ST["tables_kept_together"] += 1
            else:
                ST["tables_allowed_to_split"] += 1
            ST["tables"] += 1
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        continue

    if t.startswith('- '):
        while i < len(lines):
            cur = lines[i]
            if cur.strip().startswith('- '):
                txt = cur.strip()[2:]; j = i + 1
                while (j < len(lines) and lines[j].startswith('  ') and lines[j].strip()
                       and not lines[j].strip().startswith(('-', '|', '>', '#', '*'))):
                    txt += ' ' + lines[j].strip(); j += 1
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(5)
                p.paragraph_format.left_indent = Inches(0.3)
                add_runs(p, txt); ST["bullets"] += 1; i = j
            elif not cur.strip():
                nx = next_content_idx(i)
                if nx < len(lines) and lines[nx].strip().startswith('- '):
                    i = nx
                else:
                    break
            else:
                break
        continue

    if t.startswith('#'):
        lvl = len(t) - len(t.lstrip('#')); txt = t[lvl:].strip()
        p = doc.add_paragraph()
        if lvl == 1:
            ST["h1"] += 1
            if first_h1:
                first_h1 = False
                add_runs(p, txt, size=CFG["title_pt"], color=NAVY,
                         font=CFG["head_font"], bold_all=True)
                p.paragraph_format.space_after = Pt(2)
            else:
                IN_COVER = False          # first section heading ends the cover
                exempt = any(txt.startswith(pre) for pre
                             in CFG.get("no_page_break_h1_prefixes", []))
                if CFG["page_break_on_h1"] and not exempt:
                    p.paragraph_format.page_break_before = True
                    ST["page_breaks"] += 1
                elif exempt:
                    # Flowing section: no forced break, and NO extra space above.
                    # It takes the same heading treatment as every other section
                    # — navy heading, gold rule, identical spacing — and that
                    # treatment is what carries the structure. Do not add space
                    # here: the approved v4.5/v2.4 output was reviewed at 0 pt,
                    # and the v4.6/v2.5 baselines reproduce it exactly.
                    ST["page_breaks_exempted"] += 1
                add_runs(p, txt, size=CFG["h1_pt"], color=NAVY,
                         font=CFG["head_font"], bold_all=True)
                # Applies to broken and flowing sections alike — one source for
                # this value, so the two cannot drift apart.
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
                para_border(p, 'bottom', CFG["gold"], '10')
        elif lvl == 2:
            ST["h2"] += 1
            add_runs(p, txt, size=CFG["h2_pt"], color=NAVY,
                     font=CFG["head_font"], bold_all=True)
            p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
        else:
            ST["h3"] += 1
            add_runs(p, txt, size=CFG["h3_pt"], color=NAVY,
                     font=CFG["head_font"], bold_all=True)
            p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
        keep_with_next(p)
        i += 1; continue

    para = [t]; j = i + 1
    while (j < len(lines) and lines[j].strip()
           and not lines[j].strip().startswith(('#', '-', '|', '>', '!', '<!--'))
           and not re.match(r'^-{3,}$', lines[j].strip())):
        para.append(lines[j].strip()); j += 1
    p = doc.add_paragraph(); add_runs(p, ' '.join(para)); i = j

doc.save(OUT)
print("SAVED:", OUT)
print("PAGE_SIZE:", CFG["page_size"], f"({w_mm} x {h_mm} mm)")
print("STATS:", ST)
print("SIZE_BYTES:", os.path.getsize(OUT))


# ---------------------------------------------------------------------------
# POST-BUILD SELF-CHECK
# ---------------------------------------------------------------------------
def post_build_check(path, cfg, md_source_text):
    """Verify what can be verified without pagination.

    Pagination-dependent faults — a callout or short table actually splitting
    across a page — cannot be detected here, because page breaks are decided by
    the renderer that opens the file, not by python-docx. What CAN be checked is
    that the PROTECTIONS against them were applied. Those checks are below;
    the visual confirmation stays a human/PDF step and is recorded as such.
    """
    from docx import Document as _D
    d = _D(path)
    results, critical = [], False

    def add(ok, label, detail="", is_critical=True):
        nonlocal critical
        results.append((ok, label, detail))
        if not ok and is_critical:
            critical = True

    # 1. page size matches config
    s = d.sections[0]
    exp_w, exp_h = PAGE_SIZES_MM.get(cfg["page_size"], PAGE_SIZES_MM["A4"])
    got_w = round(s.page_width.mm); got_h = round(s.page_height.mm)
    add(abs(got_w - round(exp_w)) <= 1 and abs(got_h - round(exp_h)) <= 1,
        f"page size is {cfg['page_size']}", f"{got_w} x {got_h} mm")

    # 2. gather all text
    txt = "\n".join(p.text for p in d.paragraphs)
    for t_ in d.tables:
        for r_ in t_.rows:
            for c_ in r_.cells:
                txt += "\n" + "\n".join(p.text for p in c_.paragraphs)

    # 3. no stray Markdown emphasis markers survived conversion
    stray = txt.count("**")
    add(stray == 0, "no stray Markdown markers", f"{stray} found")

    # 4. internal markers still present while in rehearsal mode
    if cfg.get("render_internal_markers", True):
        has_boundary = "INTERNAL REHEARSAL DRAFT" in txt
        add(has_boundary, "rehearsal boundary present in output")

    # 5. callout panels carry split protection
    if cfg.get("callout_cant_split", True):
        callouts = [t_ for t_ in d.tables if len(t_.rows) == 1 and len(t_.columns) == 1]
        unprotected = [t_ for t_ in callouts
                       if not t_.rows[0]._tr.xpath('./w:trPr/w:cantSplit')]
        add(not unprotected, "callout panels protected from splitting",
            f"{len(callouts) - len(unprotected)}/{len(callouts)} protected")

    # 6. short content tables carry keep-together protection
    maxr = cfg.get("table_keep_together_max_rows", 8)
    content = [t_ for t_ in d.tables if not (len(t_.rows) == 1 and len(t_.columns) == 1)]
    short = [t_ for t_ in content if len(t_.rows) <= maxr]
    unkept = []
    for t_ in short:
        if len(t_.rows) < 2:
            continue
        first_cell = t_.rows[0].cells[0]
        if not any(p._p.xpath('./w:pPr/w:keepNext') for p in first_cell.paragraphs):
            unkept.append(t_)
    add(not unkept, "short tables protected from splitting",
        f"{len(short) - len(unkept)}/{len(short)} protected")

    # 7. output path is the configured one (guards silent fallback)
    add(os.path.abspath(path) == os.path.abspath(os.path.join(cfg["root"], cfg["out_docx"])),
        "output path matches config")

    # 8. source Markdown untouched by this build
    add(open(SRC, encoding='utf-8').read() == md_source_text,
        "source Markdown unchanged by build")

    print("\nPOST-BUILD CHECK")
    for ok, label, detail in results:
        print(f"  [{'PASS' if ok else 'FAIL'}] {label}" + (f"  ({detail})" if detail else ""))
    print("  NOT CHECKABLE HERE (needs a rendered page): actual page count, "
          "DOCX/PDF pagination match, whether any panel or table visually split.")
    if critical:
        raise SystemExit("POST-BUILD CHECK FAILED — see [FAIL] lines above.")
    print("POST-BUILD CHECK: all automated checks passed.")


if CFG.get("post_build_check", True):
    post_build_check(OUT, CFG, md_text)
